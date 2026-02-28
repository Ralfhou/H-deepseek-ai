import streamlit as st
from openai import OpenAI
from tavily import TavilyClient
import json
import concurrent.futures
import datetime
import os
from docx import Document
from io import BytesIO
from pypdf import PdfReader

# ==========================================
# 1. 基础配置
# ==========================================
st.set_page_config(page_title="DeepSeek 行动派 (Level 18)", page_icon="🦾", layout="wide")

deepseek_key = st.secrets.get("DEEPSEEK_API_KEY")
tavily_key = st.secrets.get("TAVILY_API_KEY")

if not deepseek_key or not tavily_key:
    st.error("❌ 请检查 Secrets 配置！")
    st.stop()

client = OpenAI(api_key=deepseek_key, base_url="https://api.deepseek.com")
tavily = TavilyClient(api_key=tavily_key)

if "messages" not in st.session_state: st.session_state.messages = []
if "current_report" not in st.session_state: st.session_state.current_report = "" 
if "web_context" not in st.session_state: st.session_state.web_context = "" 

# ==========================================
# 2. 机械臂：定义 AI 可以调用的本地工具
# ==========================================

# 2.1 真实的 Python 执行函数 (在你的电脑上做事)
def execute_save_to_desktop(filename, content):
    """ 将内容保存到本地桌面的专用文件夹 """
    try:
        # 自动识别 Windows 或 Mac 的桌面路径
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        agent_folder = os.path.join(desktop_path, "DeepSeek_Agent_Outputs")
        
        # 如果文件夹不存在，就自动创建一个
        os.makedirs(agent_folder, exist_ok=True)
        
        file_path = os.path.join(agent_folder, filename)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
            
        return f"✅ 执行成功！文件已保存至你的本地路径: {file_path}"
    except Exception as e:
        return f"❌ 保存失败: {str(e)}"

# 2.2 给 AI 看的技能说明书 (JSON Schema)
tools = [
    {
        "type": "function",
        "function": {
            "name": "execute_save_to_desktop",
            "description": "当用户明确要求将报告、文本或文件【保存到桌面】或【存到本地】时，调用此工具。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "文件的名称，必须包含后缀名（如 report.md, data.txt）"
                    },
                    "content": {
                        "type": "string",
                        "description": "需要保存的完整文本内容"
                    }
                },
                "required": ["filename", "content"]
            }
        }
    }
]

# ==========================================
# 3. 智能体定义 (保留之前的策划、搜索、精修)
# ==========================================

def agent_planner(query):
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    prompt = f"你是策划。今天是{today}。用户选题：'{query}'。制定3个搜索词。输出JSON: {{'queries': ['词1', '词2']}}"
    res = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "user", "content": prompt}], response_format={"type": "json_object"})
    return json.loads(res.choices[0].message.content)

def agent_searcher(queries):
    new_context = ""
    def fetch(q):
        try: return tavily.search(query=q, search_depth="advanced", max_results=2)['results']
        except: return []
    with concurrent.futures.ThreadPoolExecutor() as executor:
        for future in concurrent.futures.as_completed([executor.submit(fetch, q) for q in queries]):
            for item in future.result(): new_context += f"Source: {item['title']}\nContent: {item['content']}\n\n"
    return new_context

def agent_writer(query, context):
    prompt = f"你是撰稿人。选题：'{query}'\n资料：{context}\n请写一份Markdown格式的初稿。"
    res = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "user", "content": prompt}])
    return res.choices[0].message.content

# --- 新增: 动作决策大模型 (Action Agent) ---
def agent_action_executor(user_command, current_report):
    """ 专门负责判断是否需要调用工具，以及执行工具 """
    
    # 构建上下文对话
    messages = [
        {"role": "system", "content": "你是一个能操控用户电脑的行动派AI。如果用户要求保存文件，请务必使用你拥有的工具。"},
        {"role": "user", "content": f"这是当前的报告内容：\n{current_report}\n\n用户的指令是：'{user_command}'"}
    ]
    
    # 呼叫 DeepSeek，并把工具箱递给它
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=messages,
        tools=tools,
        tool_choice="auto" # 让 AI 自动决定是否使用工具
    )
    
    response_message = response.choices[0].message
    
    # 判断 AI 是否决定使用工具
    if response_message.tool_calls:
        tool_call = response_message.tool_calls[0]
        function_name = tool_call.function.name
        
        # 提取 AI 自动生成的参数
        function_args = json.loads(tool_call.function.arguments)
        
        # 真正地执行本地 Python 函数
        if function_name == "execute_save_to_desktop":
            result_msg = execute_save_to_desktop(
                filename=function_args.get("filename"),
                content=function_args.get("content")
            )
            return result_msg
    else:
        # 如果不需要用工具，就正常回复文本
        return response_message.content

def generate_docx(content):
    doc = Document()
    doc.add_heading('DeepSeek 研报', 0)
    for line in content.split('\n'):
        if line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        else: doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==========================================
# 4. 页面 UI
# ==========================================
with st.sidebar:
    st.header("🦾 机械臂控制台")
    if st.button("🗑️ 清空重来"):
        st.session_state.current_report = ""
        st.session_state.messages = []
        st.rerun()
    st.info("💡 提示：在获得报告后，你可以直接命令它：\n'帮我把报告保存到桌面上，命名为xxx.md'")

st.title("🦾 DeepSeek 行动派 (Tool Calling)")
st.caption("Level 18: Break out of the browser. Save files locally.")

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]): st.markdown(msg["content"])

if not st.session_state.current_report:
    if user_input := st.chat_input("请输入初始选题..."):
        st.session_state.messages.append({"role": "user", "content": user_input})
        with st.chat_message("user"): st.write(user_input)

        with st.chat_message("assistant"):
            with st.status("🚀 制作中...", expanded=True) as status:
                status.write("🧠 [策划] 制定方案...")
                plan = agent_planner(user_input)
                status.write("🌍 [猎手] 搜索中...")
                web_ctx = agent_searcher(plan['queries'])
                status.write("✍️ [主笔] 撰写中...")
                draft = agent_writer(user_input, web_ctx)
                status.update(label="✅ 初稿完成！", state="complete")

            st.markdown(draft)
            st.session_state.messages.append({"role": "assistant", "content": draft})
            st.session_state.current_report = draft
            st.rerun() 
else:
    st.download_button("📥 手动下载 (.docx)", generate_docx(st.session_state.current_report), "report.docx")
    
    with st.form("action_form"):
        feedback = st.text_area("✍️ 下达动作指令 (例如：把报告保存到我的桌面上，叫 测试.md)", placeholder="你可以让它修改文章，或者让它执行保存动作...")
        submitted = st.form_submit_button("🚀 执行指令")
        
    if submitted and feedback:
        st.session_state.messages.append({"role": "user", "content": f"【指令】{feedback}"})
        with st.chat_message("user"): st.write(f"【指令】{feedback}")
        
        with st.chat_message("assistant"):
            with st.status("🦾 正在思考并尝试执行动作...", expanded=True) as status:
                
                # 核心：调用行动派 Agent
                action_result = agent_action_executor(feedback, st.session_state.current_report)
                
                status.write(action_result)
                status.update(label="✅ 动作执行完毕", state="complete")
                
        st.session_state.messages.append({"role": "assistant", "content": action_result})
        st.rerun()